from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('Nijahomzs_Client_Action_Checklist.docx')

COLORS = {
    'navy': '0B2A4A',
    'blue': '0057B8',
    'sky': 'EAF4FF',
    'gold': 'F4B400',
    'gold_light': 'FFF4D6',
    'emerald': '0F9F6E',
    'emerald_light': 'E9F8F1',
    'red': 'D93025',
    'red_light': 'FDECEC',
    'slate': '475569',
    'slate_light': 'F1F5F9',
    'white': 'FFFFFF',
    'border': 'D7E2EF',
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, color='D7E2EF', sz='8'):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), sz)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_text(cell, text, bold=False, color='111827', size=9.5, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08


def add_run(p, text, bold=False, color='111827', size=None):
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return r


def add_badge(paragraph, text, fill, color='FFFFFF'):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(color)
    # highlight supports limited colors only, so keep text badge simple
    return run


def add_section_title(doc, title, subtitle=None, accent='blue'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(COLORS[accent])
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, subtitle, color=COLORS['slate'], size=10.5)


def add_callout(doc, title, body, fill='EAF4FF', border='0057B8'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, '12')
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=border, size=11.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body, color='1F2937', size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_action_box(doc, number, title, priority, status, why, send_items, steps, done_when, accent='blue'):
    add_section_title(doc, f'{number}. {title}', f'Priority: {priority} | Current status: {status}', accent=accent)
    add_callout(doc, 'Why this is needed', why, fill=COLORS[f'{accent}_light'] if f'{accent}_light' in COLORS else COLORS['sky'], border=COLORS[accent])

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.1)
    tbl.columns[1].width = Inches(4.6)
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(tbl.cell(0,0), COLORS['slate_light'])
    set_cell_text(tbl.cell(0,0), 'Client should send', bold=True, color=COLORS['navy'], size=9.5)
    set_cell_text(tbl.cell(0,1), '\n'.join([f'- {x}' for x in send_items]), color='1F2937', size=9.2)

    row = tbl.add_row().cells
    for cell in row:
        set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(row[0], COLORS['slate_light'])
    set_cell_text(row[0], 'How to do it', bold=True, color=COLORS['navy'], size=9.5)
    set_cell_text(row[1], '\n'.join([f'{i+1}. {x}' for i, x in enumerate(steps)]), color='1F2937', size=9.2)

    row = tbl.add_row().cells
    for cell in row:
        set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(row[0], COLORS['slate_light'])
    set_cell_text(row[0], 'Done when', bold=True, color=COLORS['navy'], size=9.5)
    set_cell_text(row[1], done_when, color='1F2937', size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string('1F2937')
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [('Title', 28, COLORS['white']), ('Heading 1', 18, COLORS['blue']), ('Heading 2', 14, COLORS['navy'])]:
        st = styles[style_name]
        st.font.name = 'Aptos Display' if style_name != 'Normal' else 'Aptos'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), st.font.name)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(4)
        r = footer.add_run('Nijahomzs Client Action Checklist | Prepared for Deployment Handoff')
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string('64748B')


def build():
    doc = Document()
    style_document(doc)

    # Cover band
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover.cell(0,0)
    set_cell_shading(cell, COLORS['navy'])
    set_cell_border(cell, COLORS['navy'], '0')
    set_cell_margins(cell, top=420, start=360, bottom=420, end=360)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, 'NIJAHOMZS', bold=True, color='FFFFFF', size=15)
    p2 = cell.add_paragraph()
    add_run(p2, 'Client Action Checklist', bold=True, color='FFFFFF', size=30)
    p3 = cell.add_paragraph()
    add_run(p3, 'Final items required from the client to unlock payments, support email, social posting, WhatsApp dashboard access, and full production QA.', color='DCEBFF', size=12)
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(10)
    add_run(p4, 'Prepared for: Nijahomzs Team', bold=True, color=COLORS['gold'], size=11)
    p5 = cell.add_paragraph()
    add_run(p5, 'Date: May 2026', color='DCEBFF', size=10)

    doc.add_paragraph()
    add_callout(doc, 'Quick Summary', 'Most core development is already built and deployed. The remaining client actions are mainly credentials, account access decisions, billing/quota setup, and approvals. Once these are provided, the developer can run final live tests and complete production handoff.', fill='E9F8F1', border=COLORS['emerald'])

    # Priority board
    add_section_title(doc, 'Client Priority Board', 'Please complete the red and amber items first. They directly block live testing and production reliability.', accent='blue')
    board = doc.add_table(rows=1, cols=5)
    board.alignment = WD_TABLE_ALIGNMENT.CENTER
    board.autofit = False
    widths = [0.55, 1.45, 1.15, 2.15, 1.2]
    headers = ['No.', 'Item', 'Priority', 'What client must provide/do', 'Status']
    for i, h in enumerate(headers):
        c = board.cell(0,i)
        c.width = Inches(widths[i])
        set_cell_shading(c, COLORS['blue'])
        set_cell_border(c, COLORS['blue'])
        set_cell_margins(c, 130, 100, 130, 100)
        set_cell_text(c, h, bold=True, color='FFFFFF', size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ('1', 'Firebase quota/billing', 'Critical', 'Resolve quota or enable billing so Firestore/API reads work reliably.', 'Blocked'),
        ('2', 'SMTP support email', 'High', 'Provide mail server credentials and support inbox address.', 'Missing'),
        ('3', 'Flutterwave + FCMB', 'High', 'Provide final keys, webhook hash, and bank settlement details.', 'Missing'),
        ('4', 'Buffer/social profiles', 'Medium', 'Provide Buffer API key and connected FB/IG/X profile IDs.', 'Missing'),
        ('5', 'WhatsApp manager access', 'Medium', 'Choose dashboard subdomain and add DNS A record.', 'Decision needed'),
        ('6', 'Business rules', 'Medium', 'Approve KYC rules, ad packages, blog sources, and content calendar.', 'Decision needed'),
        ('7', 'Mobile app accounts', 'Later', 'Prepare Google Play and Apple Developer accounts.', 'Future phase'),
    ]
    priority_colors = {'Critical': COLORS['red_light'], 'High': COLORS['gold_light'], 'Medium': COLORS['sky'], 'Later': COLORS['slate_light']}
    for row_data in rows:
        cells = board.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_border(cells[i])
            set_cell_margins(cells[i], 120, 100, 120, 100)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 2:
                set_cell_shading(cells[i], priority_colors.get(val, COLORS['slate_light']))
                set_cell_text(cells[i], val, bold=True, color=COLORS['navy'], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif i in (0,4):
                set_cell_shading(cells[i], COLORS['slate_light'])
                set_cell_text(cells[i], val, bold=(i==4), color=COLORS['navy'], size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                set_cell_text(cells[i], val, color='1F2937', size=8.5)

    doc.add_page_break()

    add_action_box(
        doc, '1', 'Fix Firebase / Firestore quota or billing', 'CRITICAL', 'Currently blocking some live APIs and workers',
        'The website, admin dashboards, support inbox, ads, content workers, onboarding worker, and public stats rely on Firestore. The VPS logs currently show quota exceeded errors, which can make completed features appear broken.',
        ['Confirm Firebase billing is enabled or quota has been increased.', 'Send a screenshot/confirmation that the Firebase project quota issue is resolved.', 'Confirm the correct Firebase project is being used for production.'],
        ['Open Firebase Console.', 'Select the Nijahomzs production project.', 'Go to Usage and billing / Firestore usage.', 'Enable billing or increase quota if prompted.', 'Wait for quota reset, then notify the developer to rerun production QA.'],
        'Developer can call Firestore-backed APIs without RESOURCE_EXHAUSTED errors, and workers stop failing due to quota.',
        accent='red'
    )

    add_action_box(
        doc, '2', 'Provide SMTP email credentials for support workflow', 'HIGH', 'Support tickets are created, but live email sending is not configured',
        'The support system is built, but email notifications and admin email replies need SMTP credentials. Without SMTP, tickets still save, but emails will not send.',
        ['SMTP_HOST, for example smtp.zoho.com or mail.yourdomain.com', 'SMTP_PORT, usually 465 or 587', 'SMTP_USER', 'SMTP_PASS or app password', 'SMTP_FROM_EMAIL, for example support@nijahomzs.com', 'SUPPORT_EMAIL_TO, the inbox where new tickets should arrive'],
        ['Decide which email provider will send support emails.', 'Create or confirm the support email account.', 'Generate an app password if the provider requires it.', 'Send the SMTP details securely to the developer.', 'Developer will add them to VPS env and test one support email.'],
        'A new contact form ticket sends an email to the support inbox, and admin replies can be sent by email.',
        accent='gold'
    )

    doc.add_page_break()

    add_action_box(
        doc, '3', 'Provide Flutterwave payment keys and FCMB settlement details', 'HIGH', 'Payment code exists, but VPS keys are missing',
        'The payment routes, webhook handler, transaction logs, zero-fee intent logging, and ad campaign payment hooks exist. Live payment testing cannot finish until Flutterwave credentials are added.',
        ['FLUTTERWAVE_PUBLIC_KEY', 'FLUTTERWAVE_SECRET_KEY', 'FLUTTERWAVE_WEBHOOK_HASH', 'Confirm test mode or live mode', 'FCMB_ACCOUNT_NUMBER', 'FCMB_ACCOUNT_NAME'],
        ['Log in to Flutterwave.', 'Open Settings / API Keys.', 'Copy the public key and secret key for the correct mode.', 'Create or copy the webhook hash/secret.', 'Confirm the settlement bank account details.', 'Send only the keys/details, not the login password, unless absolutely necessary.'],
        'Developer can initialize payment, verify payment, receive webhook, log transaction, and mark ad/promotional payments correctly.',
        accent='gold'
    )

    add_action_box(
        doc, '4', 'Provide Buffer/social media posting credentials', 'MEDIUM', 'Social posting queue is built, but Buffer is not connected',
        'The content automation system can generate blog drafts and queue social posts. To publish to Facebook, Instagram, and X through Buffer, the project needs a Buffer API key and connected profile IDs.',
        ['BUFFER_API_KEY', 'Facebook Page/Profile channel ID', 'Instagram Business/Profile channel ID', 'X/Twitter channel ID', 'Decision for TikTok and Facebook Groups workflow'],
        ['Create or open the Buffer account.', 'Connect the Facebook Page, Instagram Business account, and X account.', 'Generate/copy the Buffer API key if available for the account.', 'Provide the channel/profile IDs to the developer.', 'Confirm whether TikTok/FB Groups should be manual, Buffer-supported, or a later integration.'],
        'Approved blog posts can be queued and posted to connected social channels without exposing credentials in the frontend.',
        accent='blue'
    )

    doc.add_page_break()

    add_action_box(
        doc, '5', 'Choose and configure WhatsApp Manager dashboard access', 'MEDIUM', 'Evolution API is running and instance is open; client dashboard subdomain still needs decision/DNS',
        'Evolution API is already running on the VPS and the WhatsApp instance is open. For remote dashboard/QR access, the client should choose a secure subdomain and point DNS to the VPS.',
        ['Preferred dashboard URL: manager.nijahomzs.com or whatsapp.nijahomzs.com', 'DNS A record pointing that subdomain to 158.69.1.183', 'Preferred Basic Auth username/password for dashboard access', 'Confirm who will scan/manage the WhatsApp QR code'],
        ['Choose the subdomain.', 'In the domain DNS panel, create an A record pointing to 158.69.1.183.', 'Send the chosen subdomain and Basic Auth credentials to the developer.', 'Developer will configure Nginx + SSL and share the dashboard link.', 'Client scans QR or confirms the phone remains connected.'],
        'Dashboard opens securely over HTTPS, protected by login, and the client can manage/scan WhatsApp remotely.',
        accent='blue'
    )

    add_action_box(
        doc, '6', 'Approve business rules for KYC, ads, and content', 'MEDIUM', 'Systems exist; final rules must be client-approved',
        'The platform needs clear business rules so admins can approve users, campaigns, and content consistently. These are business decisions, not just developer tasks.',
        ['KYC rule: ID required? CAC required? both or either?', 'KYC approval team and review SLA', 'Ad package names, prices, durations, and banner sizes', 'Content topics, source links, tone, and forbidden topics', 'Who approves blogs and ad campaigns before publishing'],
        ['Review the proposed rules internally.', 'Write the final choices in a short reply or spreadsheet.', 'Provide 10-20 preferred blog/source links if available.', 'Approve monthly content volume target, for example 20-30 posts/month.', 'Developer will encode/test the final rules in admin workflows.'],
        'Admin team can approve/reject KYC, ads, and content without confusion, and reports match the client business model.',
        accent='emerald'
    )

    doc.add_page_break()

    add_action_box(
        doc, '7', 'Prepare mobile app accounts for the future phase', 'LATER', 'Native mobile apps are not part of the completed web deployment yet',
        'The web platform is mobile responsive, but native Android/iOS apps, push notifications, app store builds, and recommendation engine are a separate phase. Accounts should be prepared early to avoid launch delays.',
        ['Google Play Console account access', 'Apple Developer account access', 'App name confirmation', 'App icon/splash branding assets', 'Push notification sender/account decision'],
        ['Create/confirm Google Play Console account.', 'Create/confirm Apple Developer account.', 'Add the developer/team when mobile work starts.', 'Prepare brand assets in high resolution.', 'Confirm notification policy and user consent wording.'],
        'Mobile build can be submitted to app stores without waiting for account setup.',
        accent='slate'
    )

    add_section_title(doc, 'Secure Sharing Rules', 'Please follow these rules when sending credentials or access details.', accent='red')
    rules = [
        ('Do not send unnecessary login passwords', 'Send API keys, SMTP app passwords, webhook secrets, and IDs instead of full account logins whenever possible.'),
        ('Use a secure channel', 'Prefer a password manager share link, encrypted document, or controlled email thread. Avoid posting secrets in public groups.'),
        ('Label each item clearly', 'Example: FLUTTERWAVE_SECRET_KEY = ..., SMTP_HOST = ..., BUFFER_API_KEY = ...'),
        ('Confirm environment mode', 'Clearly say whether each key is TEST/SANDBOX or LIVE/PRODUCTION.'),
        ('Tell the developer after DNS changes', 'DNS can take time to propagate. Once updated, notify the developer to configure SSL and test.'),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(tbl.cell(0,0), COLORS['red'])
    set_cell_shading(tbl.cell(0,1), COLORS['red'])
    set_cell_text(tbl.cell(0,0), 'Rule', bold=True, color='FFFFFF', size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tbl.cell(0,1), 'Meaning', bold=True, color='FFFFFF', size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for cell in tbl.rows[0].cells:
        set_cell_border(cell, COLORS['red']); set_cell_margins(cell)
    for rule, meaning in rules:
        cells = tbl.add_row().cells
        set_cell_text(cells[0], rule, bold=True, color=COLORS['navy'], size=9.2)
        set_cell_text(cells[1], meaning, color='1F2937', size=9.2)
        for c in cells:
            set_cell_border(c); set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_section_title(doc, 'Final Handoff Checklist', 'Once the client completes the items below, the developer can run final live QA and close production handoff.', accent='emerald')
    checklist = [
        'Firebase quota/billing fixed and confirmed.',
        'SMTP support email credentials added and tested.',
        'Flutterwave keys, webhook hash, and FCMB details added and tested.',
        'Buffer API key and channel IDs added and tested.',
        'WhatsApp dashboard subdomain DNS set and SSL configured.',
        'KYC, ads, and content rules approved.',
        'Final regression QA completed across support, payments, WhatsApp, KYC, ads, blog, analytics, and listing flows.',
    ]
    for item in checklist:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, '[ ] ', bold=True, color=COLORS['emerald'], size=11)
        add_run(p, item, color='1F2937', size=10.2)

    add_footer(doc)
    doc.save(OUT)

if __name__ == '__main__':
    build()
    print(OUT.resolve())
