from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('Nijahomzs_WhatsApp_Onboarding_Explainer.docx')

COLORS = {
    'navy': '0B2A4A',
    'blue': '0057B8',
    'sky': 'EAF4FF',
    'gold': 'F4B400',
    'gold_light': 'FFF4D6',
    'green': '16A34A',
    'green_light': 'EAF8EF',
    'slate': '475569',
    'slate_light': 'F1F5F9',
    'white': 'FFFFFF',
    'border': 'D7E2EF',
}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def border(cell, color='D7E2EF', size='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), size)
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)


def margins(cell, top=110, start=130, bottom=110, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for k, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{k}'))
        if node is None:
            node = OxmlElement(f'w:{k}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_text(cell, text, bold=False, color='1F2937', size=9.4, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def run(p, text, bold=False, color='1F2937', size=10):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def section_title(doc, title, color='0057B8'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    run(p, title, bold=True, color=color, size=14)


def flow_box(cell, title, body, fill, accent):
    shade(cell, fill)
    border(cell, accent, '10')
    margins(cell, 120, 130, 120, 130)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run(p, title, bold=True, color=accent, size=10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    run(p2, body, color='1F2937', size=8.5)


def bullet(doc, text, color='1F2937'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(1.8)
    run(p, '• ', bold=True, color=COLORS['green'], size=9.6)
    run(p, text, color=color, size=9.3)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.35)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Aptos'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string('1F2937')
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    # Header
    header = doc.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcell = header.cell(0, 0)
    shade(hcell, COLORS['navy'])
    border(hcell, COLORS['navy'], '0')
    margins(hcell, 220, 240, 220, 240)
    p = hcell.paragraphs[0]
    run(p, 'NIJAHOMZS WHATSAPP ONBOARDING', bold=True, color='FFFFFF', size=10.5)
    p2 = hcell.add_paragraph()
    run(p2, 'How Scraped Property Agents Receive Claim Messages', bold=True, color='FFFFFF', size=20)
    p3 = hcell.add_paragraph()
    run(p3, 'A simple explanation of the automated workflow, the service used, and the business benefit.', color='DCEBFF', size=10.2)

    # Simple summary callout
    summary = doc.add_table(rows=1, cols=1)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = summary.cell(0, 0)
    shade(c, COLORS['green_light'])
    border(c, COLORS['green'], '12')
    margins(c, 130, 160, 130, 160)
    p = c.paragraphs[0]
    run(p, 'In one line: ', bold=True, color=COLORS['green'], size=10.2)
    run(p, 'When Nijahomzs imports a property advert, the system can automatically send the agent a WhatsApp message with a secure claim link so they can sign up, claim, and manage their listing.', color='1F2937', size=10.2)

    services = doc.add_table(rows=1, cols=3)
    services.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ('Evolution API', 'Self-hosted WhatsApp service on the VPS. It sends messages using the connected Nijahomzs WhatsApp number.', COLORS['sky'], COLORS['blue']),
        ('Firestore Queue', 'Stores pending outreach jobs, claim tokens, send status, retries, and opt-outs.', COLORS['slate_light'], COLORS['navy']),
        ('Node.js Worker', 'Runs in the background, sends one message at a time, and waits between messages to reduce WhatsApp risk.', COLORS['gold_light'], COLORS['gold']),
    ]
    for i, (title, body, fill, accent) in enumerate(items):
        flow_box(services.cell(0, i), title, body, fill, accent)

    flow = doc.add_table(rows=1, cols=5)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    steps = [
        ('1. Import', 'A property advert is scraped/imported into Nijahomzs.'),
        ('2. Queue', 'The system finds the agent phone number and creates a pending WhatsApp job.'),
        ('3. Claim Link', 'A secure 14-day claim token is generated for that advert.'),
        ('4. WhatsApp', 'Evolution API sends the agent a fixed onboarding message with the claim link.'),
        ('5. Claim', 'The agent signs up/logs in and the advert becomes linked to their account.'),
    ]
    for i, (title, body) in enumerate(steps):
        flow_box(flow.cell(0, i), title, body, 'FFFFFF', COLORS['blue'] if i % 2 == 0 else COLORS['green'])

    msg = doc.add_table(rows=1, cols=1)
    msg.alignment = WD_TABLE_ALIGNMENT.CENTER
    m = msg.cell(0, 0)
    shade(m, COLORS['gold_light'])
    border(m, COLORS['gold'], '12')
    margins(m, 130, 170, 130, 170)
    p = m.paragraphs[0]
    run(p, 'What Message Does The Agent Receive? ', bold=True, color=COLORS['gold'], size=10.5)
    run(p, 'Example: ', bold=True, color=COLORS['gold'], size=9.8)
    run(p, '"Your property advert is now live on Nijahomzs. To take full control, edit the details, and keep the listing free, claim it here: [secure claim link]. Reply STOP to opt-out."', color='1F2937', size=9.5)

    benefits = doc.add_table(rows=1, cols=2)
    benefits.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = benefits.cell(0, 0)
    right = benefits.cell(0, 1)
    for cell in [left, right]:
        shade(cell, 'FFFFFF')
        border(cell, COLORS['border'])
        margins(cell, 120, 160, 120, 160)
    set_text(left, 'Why This Is Useful For Nijahomzs\n\nBusiness Benefits\n- Converts imported adverts into registered users\n- Brings agents into the platform automatically\n- Encourages agents to update and improve listings\n- Helps build relationships with property agents\n- Saves manual outreach time', bold=False, color='1F2937', size=9.1)
    set_text(right, 'Safety Controls\n\n- Secure claim token, not a public takeover link\n- Token expires after 14 days\n- Duplicate messages are avoided\n- STOP replies are saved as opt-outs\n- Messages are throttled with delays and daily/hourly caps', bold=False, color='1F2937', size=9.1)

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    n = note.cell(0, 0)
    shade(n, COLORS['slate_light'])
    border(n, COLORS['navy'], '10')
    margins(n, 100, 150, 100, 150)
    p = n.paragraphs[0]
    run(p, 'Important Note: ', bold=True, color=COLORS['navy'], size=9.4)
    run(p, 'For this to work reliably, the WhatsApp number must remain connected in Evolution API, Firestore quota must be healthy, and the background worker must stay online on the VPS.', bold=True, color=COLORS['navy'], size=9.2)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(footer, 'Nijahomzs Automated WhatsApp Onboarding | Client Explanation', color='64748B', size=7.8)

    doc.save(OUT)

if __name__ == '__main__':
    build()
    print(OUT.resolve())
